import time
import spacy
import pathlib
import os
import docx
import PyPDF2
import csv
from typing import Optional
from pydantic import BaseModel
import logging
from dotenv import load_dotenv
import requests
import glob
import json
import aspose.words as aw


GROUP_ID = 'new_group32939'
JOB_ID = '105'

CANDIDATE_CSV_PATH = 'candidate.csv'
CANDIDATE_NOTES_CSV_PATH = 'candidatenote.csv'
RESUME_CSV_PATH = 'resume.csv'
RESUME_FOLDER_PATH = '/Volumes/T7/sbt_files/unzip/resumes/'
RECRUITERS_CSV_PATH = 'recruiter.csv'
LOGGER_PATH = 'receipt.log'
ENV_PATH = '../../.env'
load_dotenv(dotenv_path=ENV_PATH)

BOARD_ID = 4846750007


# Logging formation
logger = logging.getLogger('logger')

# Set logger level
logger.setLevel(logging.DEBUG)

# create handler
handler = logging.FileHandler(LOGGER_PATH)

# set handler level
handler.setLevel(logging.DEBUG)

# create formatter
formatter = logging.Formatter(
    '%(asctime)s - %(name)s - %(levelname)s - %(message)s')

# set formatter
handler.setFormatter(formatter)

# add handler to logger
logger.addHandler(handler)


class CandidateData(BaseModel):
    hasResume: Optional[bool] = None
    location: Optional[str] = None
    LA_area: Optional[str] = None
    name: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    date: Optional[str] = None
    pagelink: Optional[str] = None
    source: Optional[str] = None
    license_cert: Optional[str] = None
    license_expiration: Optional[str] = None
    title: Optional[str] = None
    monday_id: Optional[str] = None
    notes_id: Optional[str] = None
    notes: Optional[str] = None
    hasNotes: Optional[bool] = None
    resume_file: Optional[str] = None
    jobdiva_id: Optional[str] = None


def get_raw_text(file_path: str) -> str:
    """
    Extracts raw text from a document file. Supports .docx, .pdf, and .rtf formats.

    Args:
        file_path (str): The path to the document file.

    Returns:
        str: The extracted text from the document.

    Raises:
        ValueError: If the file format is unsupported.

    Notes:
        - For .docx files, extracts text from paragraphs.
        - For .pdf files, extracts text from all pages.
        - For .rtf files, converts the RTF format to pdf.
    """
    # Define text string and get file extension
    text = ""
    file_extension = pathlib.Path(file_path).suffix.lower().strip()
    print(file_extension)
    # If .docx file
    if file_extension == '.docx':
        # Load the .docx file and join paragraphs
        doc = docx.Document(file_path)
        text = ' '.join([para.text for para in doc.paragraphs])

    # If .pdf file
    elif file_extension == '.pdf':
        with open(file_path, 'rb') as f:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(f)

            # Get the number of pages in the PDF file
            num_pages = len(pdf_reader.pages)

            # Loop through all the pages and extract the text
            for page in range(num_pages):
                # Get the page object
                pdf_page = pdf_reader.pages[page]
                # Extract the text from the page
                page_text = pdf_page.extract_text()
                if page_text:
                    text += page_text

    # If .rtf file
    elif file_extension == '.rtf':
        with open(file_path, 'rb') as f:
            # Parse the .rtf content
            text = file_path.read()

    # Unsupported file format
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

    return text


def retriever_name_from_text(text: str):
    """
    Uses pre-trained NLP to perform NER on given text, returns the first such match of Label: PERSON
    Args:
        text (str): The string to be parsed

    Returns:
        str: The name which Label:PERSON was matched to

    Notes:
        - Looking into solutions in the case that matching a Label:PERSON fails
        - Will likely return "" in the case that nobody is matched, meaning that the
    """

    # Load the spaCy English language model
    nlp = spacy.load("en_core_web_trf")

    # Process the text with the spaCy model
    doc = nlp(text)

    # Extract named entities from the processed text
    for ent in doc.ents:
        print(f"Entity: {ent.text}, Label: {ent.label_}")


def load_candidate_data_from_csv(row: dict) -> CandidateData:
    """
    Step 2: Load candidate data from the CSV row into the CandidateData object.
    """
    candidate = CandidateData(
        name=f"{row['FIRSTNAME']} {row['LASTNAME']}",
        email=row['EMAIL'],
        location=f"{row['ADDRESS1']} {row['ZIPCODE']} {row['CITY']} {row['STATE']}",
        phone=row['CELLPHONE'],
        jobdiva_id=row['ID']
    )
    return candidate


def check_monday_for_candidate(name: str) -> bool:
    """
    Step 3: Check if the candidate already exists on Monday.com.

    Placeholder for Monday.com API call:
    If the candidate exists, return True; otherwise, return False.
    """
    # TODO: Implement Monday.com API call to check for candidate existence
    logger.info(f"{name} - Checking if already exists on Monday")

    # Monday Code
    apiKey = os.environ['MONDAY_API_KEY']
    headers = {
        "Authorization": apiKey,
        "API-version": "2024-10"
    }
    url = "https://api.monday.com/v2"

    # Generate Query
    # TODO Add code that defaults to fail if we cannot find the proper group id
    q = f"""
        {{
        items_page_by_column_values (limit: 50, board_id: {BOARD_ID}, columns: [{{column_id: "name", column_values: [\"{name}\"]}}]) {{
            items {{
                id
                name
            }}
        }}
        }}
        """

    # Send request
    response = requests.post(url=url, headers=headers, json={'query': q})

    # Handle response
    if response.status_code == 200:
        # Extract data
        response_data = response.json()
        response_list = response_data['data']['items_page_by_column_values']['items']
        # Log
        # logger.info(f"{name} - Successfully got info from monday")

        if (len(response_list) > 0):
            logger.info(
                f"-------> {name} already exists on Monday.com, skipping.")
            return True
        else:
            logger.info(
                f"{name} - Candidate not found on Monday.com, continuing process.")
            return False
    else:
        # Handle get info fail
        logger.info(
            f"{name} Failed to get Info from Monday, continuing process.")
        return False


def search_resume_by_global_id(candidate_id: str) -> Optional[str]:
    """
    Step 4: Search resume.csv for a globalID associated with the candidateID.
    """

    def convert_rtf_to_pdf(rtf_file: str) -> str:
        """
        Converts an RTF file to a PDF.
        This is a placeholder for the actual conversion logic.
        Returns the new PDF file name.
        """
        logger.info(f'Found RTF file instead of PDF, converting file to PDF')
        doc = aw.Document(rtf_file)

        pdf_file = rtf_file.replace('.rtf', '.pdf')

        logger.info(f"Converting {rtf_file} to {pdf_file} ")

        doc.save(pdf_file)

        return pdf_file

    def chooseReturnFile(files: list) -> Optional[str]:
        """
        Helper Function:
        Given a list of file names, this function returns a suitable file name.

        If a file with a .pdf extension exists, return that file.
        If not, and a file with an .rtf extension exists, convert it to a .pdf and return the new file name.
        """
        # Check for a PDF file
        for file in files:
            if file.lower().endswith('.pdf'):
                return file

        # Check for an RTF file and convert it to a PDF
        for file in files:
            if file.lower().endswith('.rtf'):
                converted_file = convert_rtf_to_pdf(file)
                return converted_file

        # Return None if no suitable file is found
        return None

    try:
        # Search in resume.csv
        with open(RESUME_CSV_PATH, mode='r') as file:
            reader = csv.DictReader(file)
            for row in reader:
                if row['CANDIDATEID'] == candidate_id:
                    # Extract the GLOBAL_ID from the matching row
                    global_id = row['GLOBAL_ID']

                    logger.info(
                        f'Found Resume Identification number: {global_id}')
                    # Find matching files
                    matching_files = glob.glob(os.path.join(
                        RESUME_FOLDER_PATH, f"{global_id}*"))

                    # Log
                    if len(matching_files) > 0:
                        logger.info(f'Found Resume file(s): {matching_files}')
                    else:
                        logger.info(f'Did not find resume files')
                        return None

                    # Return the file that we want from the list
                    chosen_file = chooseReturnFile(matching_files)
                    logger.info(f'Found file: {chosen_file}')

                    return chosen_file if chosen_file else None
        return None

    except FileNotFoundError:
        print(f"{RESUME_CSV_PATH} not found.")
        return None


def search_notes_by_jobdiva_id(jobdiva_id: str) -> Optional[str]:
    """
    Search candidatenote.csv for all rows matching the jobdiva_id and return
    a formatted string with Recruiter, Date, and Note for each matching entry.
    """
    try:
        # Load recruiters data into a dictionary for quick lookup by RECRUITERID
        recruiter_dict = {}
        with open(RECRUITERS_CSV_PATH, mode='r') as recruiters_file:
            recruiter_reader = csv.DictReader(recruiters_file)
            for recruiter in recruiter_reader:
                recruiter_dict[recruiter['ID']
                               ] = f"{recruiter['FIRSTNAME']} {recruiter['LASTNAME']}"

        # Prepare a list to hold all matching notes
        notes_list: list[str] = []

        # Search the candidatenote.csv for notes corresponding to the jobdiva_id
        with open(CANDIDATE_NOTES_CSV_PATH, mode='r') as notes_file:
            notes_reader = csv.DictReader(notes_file)
            for row in notes_reader:
                if row['CANDIDATEID'] == jobdiva_id:
                    # Get recruiter name using RECRUITERID from the current row
                    recruiter_name = recruiter_dict.get(
                        row['RECRUITERID'], "Unknown Recruiter")
                    date_created = row['DATECREATED']
                    note = row['NOTE']

                    # Log
                    logger.info(f"Found candidate notes by {recruiter_name}")

                    # Append formatted note details to the list
                    notes_list.append(
                        f"Recruiter: {recruiter_name}\nDate: {date_created}\nNote: {note}\n")

        # If there are any matching notes, return them as a single formatted string
        if notes_list:
            # logger.info("\n\n".join(notes_list))
            return "\n\n".join(notes_list)
        else:
            # Return None if no matching jobdiva_id is found
            logger.info(f"Did not find any candidate notes")
            return None

    except FileNotFoundError as e:
        print(f"File not found: {e}")
        return None


async def createMondayItem(data: CandidateData, group: str, job: str):
    """
    Principal function to upload info to Monday DB,
    -> returns ID of newly created item (if created) 
    """

    # Monday Code
    apiKey = os.environ['MONDAY_API_KEY']
    apiUrl = "https://api.monday.com/v2"
    headers = {"Authorization": apiKey, "API-version": "2023-04"}

    # Assign Column values if found
    column_values = {
        "label": JOB_ID,  # Occupation
        "status4": data.source,  # Source
        "text8": data.location,  # Location
        "phone": data.phone,  # Phone
        "email": {"email": data.email, "text": data.email},  # Email
        "text": data.license_cert  # License/Certification (null OK)
    }

    j = json.dumps(column_values)
    j = j.replace('"', '\\\"')

    # Generate mutation
    createItemMutation = f'''mutation {{
        create_item (
            board_id: {BOARD_ID} 
            group_id: "{GROUP_ID}" 
            item_name: "{data.name}" 
            column_values: "{j}") {{id}}
            }}'''

    # Create query
    q = {"query": createItemMutation}

    # Send query
    response = requests.post(url=apiUrl, json=q, headers=headers)

    # Handle response
    if response.status_code == 200:
        response_data = response.json()
        new_id = response_data['data']['create_item']['id']
        logger.info(
            f"{data.name} - Successfully Added Candidate to Monday at: {new_id}")
        return new_id
    else:
        logger.info(f"{data.name} - Failed to upload to Monday")
        logger.info(response.json())
        return None


def run(start_idx: int, end_idx: int):
    """
    Driver code for running program on resumes

    Open candidate.csv and step through denoted range of candidates (0 - 100)
    Create a passed down candidateData object
    For each row, complete steps 1-5:

    - Step 1:
        Pull info from candidate.csv and update candidateData object:
            Name (FIRSTNAME, " ", LASTNAME)
            Email (EMAIL)
            Address (ADDRESS1, " ", ZIPCODE, " ", CITY, " ", STATE )
            Phone Number (CELLPHONE)
            jobdiva_id (ID)

    - Step 2:
        Ask Monday.com if this candidate name already exists:
            if it does -> exit
            if it does not -> proceed

    - Step 3:
        Search resume.csv for global_ID in row matching jobdiva_id
        Search folder for resume file with matching global_ID

    - Step 4:
        Search candidatenotes.csv for row matching jobdiva_id
        Search recruiter.csv to find the matching recruiter_id and denote the person who made the note
        TODO - Write code to format row data for legible note

    - Step 5:
        Upload candidate data from candidateData object to Monday.com

    Notes:
        - Ignoring dates for now
        - Group labels are impossible to determine through the given data, will have to do by hand later
        - Job labels are mostly listed for candidates we already have, and are mostly missing other than that,
            will also have to complete by hand after uploading to monday
    """

    try:
        logger.info(
            f'------------------ STARTING CANDIDATE ADD: {start_idx} - {end_idx} ------------------')
        with open(CANDIDATE_CSV_PATH, mode='r') as file:
            reader = csv.DictReader(file)
            for idx, row in enumerate(reader):
                # Process correct rows
                if idx < start_idx:  # Skip rows before the start index
                    continue
                if idx >= end_idx:  # Stop processing once end index is reached
                    break

                logger.info(
                    f'---------------------------------------------#{idx}---------------------------------------------')

                # Step 1: Load candidate data
                candidate_data = load_candidate_data_from_csv(row)
                logger.info(f'{candidate_data.name} - Retrieved CSV data')
                logger.info(
                    f'{candidate_data.name} - Candidate ID - {candidate_data.jobdiva_id}')

                # Check for malformed data and skip
                if candidate_data.name.lower() == "style":
                    logger.info(
                        f'{candidate_data.name} - Detected Malformed entry, skipping')
                    continue

                # Step 2: Check if candidate exists on Monday.com
                if check_monday_for_candidate(candidate_data.name):
                    continue

                # Step 3: Search resume
                resume_path = search_resume_by_global_id(
                    candidate_data.jobdiva_id)
                candidate_data.hasResume = resume_path is not None
                candidate_data.resume_file = resume_path

                # Step 4: Search for candidate notes
                notes = search_notes_by_jobdiva_id(candidate_data.jobdiva_id)
                candidate_data.hasNotes = notes is not None
                candidate_data.notes = notes

                # Step 5: Upload candidate data to Monday.com
                # TODO: Implement the upload functionality using Monday.com API
                createMondayItem()

                # Log
                logger.info(
                    f"{candidate_data.name} - Finished Processing candidate")

                # Avoid Monday rate limit
                time.sleep(0.25)

    except FileNotFoundError:
        print(f"{CANDIDATE_CSV_PATH} not found.")


# Test Group
# Nurse Job

run(6199, 6300)
