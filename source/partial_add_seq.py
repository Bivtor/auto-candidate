from auto_candidate import *
from update_monday import *
import pathlib
import asyncio
import requests
import time
import docx
import PyPDF2
import time


async def setIsUpdating(b: bool) -> bool:
    with open(WORKING_PATH, "r") as f:
        data = json.load(f)
    data['isUpdating'] = b
    with open(WORKING_PATH, "w") as outfile:
        logger.info(f"--- updating->{b} ---")  # Log working status change
        json.dump(data, outfile)


async def checkIsUpdating() -> bool:
    f = open(WORKING_PATH)
    json_check = json.load(f)
    f.close()
    return json_check['isUpdating']


async def beginHalfAdd(candidateData: candidateData):
    """
    Creates an entry with the given information, 
    returns the id of that entry so that resume and Notes can be uploaded to it
    """

    # Clean up Date TODO fix?
    if not (candidateData.date is None or candidateData.date == 'None'):
        candidateData.date = cleanupdate(
            candidateData.date, candidateData.source)

    # Send to database / return ID of this candidate?
    return await createMondayItem(candidateData)


async def finishHalfAdd(candidateData: candidateData):
    """
    Download Resume
    Update item details
    Upload Resume
    Upload Notes Sheet
    """

    # Read sheet positional settings data
    f = open(SETTINGS_PATH)
    setting_data = json.load(f)

    # Call function to finish downloading resume and update Candidate object if there is one
    if candidateData.hasResume:
        await safeCandidateUpdate(candidateData, setting_data)

    # Push item updates to Database if changes are found
    if candidateData.monday_id != None and candidateData.hasResume is True:
        await updateMondayItem(candidateData)
    else:
        logger.info(
            f"{candidateData.name} - Reportedly DID NOT upload a Resume")

    # Upload Candidate resume (assuming resume is located at newest glob)
    if candidateData.hasResume:
        await uploadCandidateResume(candidateData)
    else:
        logger.info(
            f"{candidateData.name} - Candidate Reportedly did not have a Resume")

    # Create Monday Questionairre document and
    # Update Monday LA Area Location Field
    if candidateData.monday_id != None:
        await createQuestionDocument(candidateData)
        await updateQuestionDocument(candidateData)
        updateLA_Area(name=candidateData.name,
                      monday_id=candidateData.monday_id,
                      loc=candidateData.location)

    return


async def safeCandidateUpdate(candidateData: candidateData, setting_data: dict) -> str:
    """
    Safely download and parse the resume to ensure correctness 
    1. Establish listener on download folder
    2. Open download link
    3. Upon completion, glob newest file
    4. Return stringified resume 
    """

    # 1: Open download link (Unix)

    # OPEN chrome - formatting done in paths.py
    cmd = OPEN_CMD_LINK.format(candidateData.resume_download_link)
    if os.system(cmd) != 0:
        logger.error(f"{candidateData.name} - Failed to open Resume Link")
        return
    logger.info(f"{candidateData.name} - Successfully Opened resume link")

    # 2: Wait until download finishes
    time.sleep(5)
    # TODO fix this library thing?

    # 3: Get newest file
    list_of_files = glob.glob(DOWNLOAD_PATH)  # Add back  + "/*" for Mac
    newest_file_path = max(list_of_files, key=os.path.getctime)
    print(newest_file_path)
    logger.info(f"{candidateData.name} - Resume Path: {newest_file_path}")

    # 4: Get raw text from resume
    resume_text = get_raw_text(newest_file_path)

    # 5: Update candidate Phone # based on regex
    parse_resume(candidateData, resume_text)

    # 6: Update title field (unneccessary with monday)
    category = setting_data['occupation']
    candidateData.title = "{} {} ({})".format(candidateData.name,
                                              category, candidateData.location)

    return


def get_raw_text(directory: str) -> str:
    # Define text string and get file extension
    text = ""
    file_extension = pathlib.Path(directory).suffix.strip()

    # If .docx
    if file_extension == '.docx':
        doc = docx.Document(directory)
        text = '\s'.join([para.text for para in doc.paragraphs])
   # Else assume .pdf
    else:
        with open(directory, 'rb') as f:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(f)

            # Get the number of pages in the PDF file
            num_pages = len(pdf_reader.pages)

            # Loop through all the pages and extract the text
            for page in range(num_pages):
                # Get the page object
                pdf_page = pdf_reader.pages[page]
                # Extract the text from the page
                page_text = pdf_page.extract_text()
                # Add the page text to the overall text variable
                text += page_text
    return text


async def main():
    finishHalfAdd()


if __name__ == '__main__':
    asyncio.run(main())
